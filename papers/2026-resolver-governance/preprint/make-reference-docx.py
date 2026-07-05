#!/usr/bin/env python3
"""make-reference-docx.py — generate the pandoc DOCX reference doc used by
build-submission.sh, with a centered "Page X of Y" footer (Word PAGE/NUMPAGES
fields). Pandoc copies this reference doc's footer into the built .docx, giving
the Word review copy live page numbers.

Reproducible tooling: run this to (re)create submission/reference.docx. The
generated .docx is committed as a build asset; this script is its recipe.

Usage:  python3 scripts/make-reference-docx.py
Output: submission/reference.docx
"""
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE = Path(__file__).resolve().parent.parent  # the v2 pilot dir
OUT = HERE / "submission" / "reference.docx"


def _field(run, instr):
    """Append a Word field (e.g. PAGE, NUMPAGES) to a run as begin/instr/end."""
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for el in (begin, instr_el, end):
        run._r.append(el)


def main():
    # Start from pandoc's own default reference doc so all base styles match.
    base = OUT.parent / ".reference-base.docx"
    with open(base, "wb") as fh:
        subprocess.run(
            ["pandoc", "-o", "/dev/stdout", "--print-default-data-file", "reference.docx"],
            check=True, stdout=fh,
        )

    doc = Document(str(base))
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False

    p = footer.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r1 = p.add_run("Page ")
    r2 = p.add_run()
    _field(r2, "PAGE")
    p.add_run(" of ")
    r3 = p.add_run()
    _field(r3, "NUMPAGES")

    doc.save(str(OUT))
    base.unlink(missing_ok=True)
    print(f"[make-reference-docx] wrote {OUT.relative_to(HERE)}")


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:  # noqa: BLE001
        print(f"[make-reference-docx] failed: {exc}", file=sys.stderr)
        sys.exit(1)
